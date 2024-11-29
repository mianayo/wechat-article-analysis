import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime

class ReportGenerator:
    def __init__(self):
        self.report_dir = 'output/reports'
        os.makedirs(self.report_dir, exist_ok=True)
        
    def generate_report(self, analysis_results, output_dir):
        """生成分析报告"""
        # 创建Word文档
        doc = Document()
        
        # 添加报告标题
        title = doc.add_heading('君乐宝公众号文章分析报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告生成时间
        doc.add_paragraph(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # 1. 总体统计
        doc.add_heading('1. 总体统计', level=1)
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        # 填充统计数据
        stats_data = [
            ('分析文章总数', analysis_results['total_articles']),
            ('平均字数', f"{analysis_results['overall_stats']['avg_word_count']:.0f}"),
            ('平均情感得分', f"{analysis_results['overall_stats']['avg_sentiment_score']:.2f}"),
            ('主题数量', len(analysis_results['overall_stats']['topic_distribution']))
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = str(value)
        
        # 2. 主题分布分析
        doc.add_heading('2. 主题分布分析', level=1)
        self._generate_topic_distribution_chart(
            analysis_results['overall_stats']['topic_distribution'],
            output_dir
        )
        doc.add_picture(os.path.join(output_dir, 'topic_distribution.png'), width=Inches(6))
        
        # 3. 文章详细分析
        doc.add_heading('3. 文章详细分析', level=1)
        for article in analysis_results['article_analysis']:
            doc.add_heading(article['title'], level=2)
            
            # 添加文章分析表格
            article_table = doc.add_table(rows=5, cols=2)
            article_table.style = 'Table Grid'
            
            article_data = [
                ('字数', article['word_count']),
                ('情感得分', f"{article['sentiment_score']:.2f}"),
                ('主题分类', article['topic']),
                ('可读性指数', f"{article['readability']['complexity_score']:.2f}"),
                ('关键词', '、'.join([word for word, _ in article['keywords']['textrank'][:5]]))
            ]
            
            for i, (key, value) in enumerate(article_data):
                article_table.cell(i, 0).text = key
                article_table.cell(i, 1).text = str(value)
            
            doc.add_paragraph('')  # 添加空行作为分隔
        
        # 保存报告
        report_path = os.path.join(output_dir, 'analysis_report.docx')
        doc.save(report_path)
        
    def _generate_topic_distribution_chart(self, topic_distribution, output_dir):
        """生成主题分布图表"""
        plt.figure(figsize=(10, 6))
        sns.barplot(
            x=list(topic_distribution.values()),
            y=list(topic_distribution.keys())
        )
        plt.title('文章主题分布')
        plt.xlabel('文章数量')
        plt.ylabel('主题')
        
        # 保存图表
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'topic_distribution.png'))
        plt.close()
